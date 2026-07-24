import io
import json
import logging
import zipfile
from docx import Document

logger = logging.getLogger("ai_ocr_system.exporter")

def generate_docx_bytes(results: list[dict] | dict) -> bytes:
    """
    Generates a DOCX file from OCR results.
    Works for both image OCR results (dict) and PDF page results (list of dicts).
    """
    doc = Document()
    doc.add_heading("Hasil Ekstraksi OCR", 0)
    
    if isinstance(results, list):
        # PDF OCR results
        for page_data in results:
            page_num = page_data.get("page", 1)
            doc.add_heading(f"Halaman {page_num}", level=1)
            
            text = page_data.get("text", "")
            if text:
                for line in text.split("\n"):
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph("[Tidak ada teks terdeteksi di halaman ini]")
    else:
        # Image OCR results (single page)
        doc.add_heading("Hasil Ekstraksi Gambar", level=1)
        text_regions = results.get("text_regions", [])
        if text_regions:
            for region in text_regions:
                doc.add_paragraph(region.get("text", ""))
        else:
            doc.add_paragraph("[Tidak ada teks terdeteksi]")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()


def generate_zip_bytes(filename: str, text_content: str, json_data: dict, docx_bytes: bytes, pdf_bytes: bytes | None = None) -> bytes:
    """
    Creates a ZIP archive containing TXT, JSON, DOCX, and optionally the Searchable PDF.
    """
    zip_buffer = io.BytesIO()
    base_name = filename.rsplit(".", 1)[0]
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        # 1. Add TXT file
        zip_file.writestr(f"{base_name}.txt", text_content)
        
        # 2. Add JSON file
        json_str = json.dumps(json_data, indent=2, ensure_ascii=False)
        zip_file.writestr(f"{base_name}.json", json_str)
        
        # 3. Add DOCX file
        zip_file.writestr(f"{base_name}.docx", docx_bytes)
        
        # 4. Add Searchable PDF if available
        if pdf_bytes:
            zip_file.writestr(f"{base_name}_searchable.pdf", pdf_bytes)
            
    return zip_buffer.getvalue()